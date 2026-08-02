"""
FILE NAME: docx_exporter.py

PURPOSE:
Export itinerary/costing to an A4 DOCX document.

DEPENDENCIES:
python-docx

LAST UPDATED:
2026-08-02
"""

from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT


class DOCXExporter:

    def export(self, filename, text):

        doc = Document()

        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)

        for line in str(text).split("\n"):
            stripped = line.strip()

            if not stripped:
                continue

            if stripped.isupper() and len(stripped) > 3:
                doc.add_heading(stripped, level=1)
            else:
                doc.add_paragraph(stripped)

        doc.save(filename)

        return filename
