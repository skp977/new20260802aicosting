"""
============================================================
FILE NAME
document_parser.py

PURPOSE
Parse DOCX / PDF / TXT / RTF documents into raw text.

INPUT
File path or bytes

OUTPUT
{"text": ..., "source": ..., "type": "document"}

DEPENDENCIES
pypdf, python-docx, striprtf

LAST UPDATED
2026-08-02
============================================================
"""

from pathlib import Path


class DocumentParser:

    def parse(self, source):

        path = Path(str(source))

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            text = self._parse_pdf(path)
        elif suffix == ".docx":
            text = self._parse_docx(path)
        elif suffix == ".rtf":
            text = self._parse_rtf(path)
        else:
            text = self._parse_txt(path)

        return {
            "text": text,
            "source": str(path),
            "type": "document",
            "extension": suffix
        }

    def _parse_pdf(self, path):
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = []

            for page in reader.pages:
                pages.append(
                    page.extract_text() or ""
                )

            return "\n".join(pages)
        except Exception:
            return self._read_fallback(path)

    def _parse_docx(self, path):
        try:
            from docx import Document

            document = Document(str(path))

            paragraphs = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.text
            ]

            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    paragraphs.append(" | ".join(cells))

            return "\n".join(paragraphs)
        except Exception:
            return self._read_fallback(path)

    def _parse_rtf(self, path):
        try:
            from striprtf.striprtf import rtf_to_text

            raw = path.read_text(encoding="utf-8", errors="ignore")

            return rtf_to_text(raw)
        except Exception:
            return self._read_fallback(path)

    def _parse_txt(self, path):
        return self._read_fallback(path)

    def _read_fallback(self, path):
        try:
            return path.read_text(
                encoding="utf-8",
                errors="ignore"
            )
        except Exception:
            return ""
